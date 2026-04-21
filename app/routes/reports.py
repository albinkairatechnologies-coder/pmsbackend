from flask import Blueprint, request, jsonify, make_response
from flask_jwt_extended import jwt_required, get_jwt_identity, get_jwt
from app.models.other import WorkLog, CompanySettings
from app.models.user import User
import io, csv, os
from datetime import datetime

reports_bp = Blueprint('reports', __name__)

ADMIN_ROLES = ['admin', 'team_lead', 'crm_head', 'marketing_head']


def _serialize(rows):
    """Convert timedelta/date objects to strings for JSON."""
    result = []
    for row in rows:
        clean = {}
        for k, v in row.items():
            if hasattr(v, 'total_seconds'):  # timedelta (TIME columns)
                total = int(v.total_seconds())
                clean[k] = f"{total // 3600:02d}:{(total % 3600) // 60:02d}"
            elif hasattr(v, 'isoformat'):
                clean[k] = v.isoformat()
            else:
                clean[k] = v
        result.append(clean)
    return result


def _get_pdf_fonts():
    """Register a Unicode font for ₹ support. Returns (normal, bold) font names."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os
    # Try Windows Arial first, then Linux paths
    candidates = [
        ('C:/Windows/Fonts/arial.ttf',   'C:/Windows/Fonts/arialbd.ttf'),
        ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
         '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'),
        ('/usr/share/fonts/dejavu/DejaVuSans.ttf',
         '/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf'),
    ]
    for normal_path, bold_path in candidates:
        if os.path.exists(normal_path):
            try:
                pdfmetrics.registerFont(TTFont('UniFont',      normal_path))
                pdfmetrics.registerFont(TTFont('UniFont-Bold', bold_path if os.path.exists(bold_path) else normal_path))
                return 'UniFont', 'UniFont-Bold'
            except Exception:
                continue
    return 'Helvetica', 'Helvetica-Bold'


@reports_bp.route('/reports/client-summary', methods=['GET'])
@jwt_required()
def report_client_summary():
    claims = get_jwt()
    if claims['role'] not in ADMIN_ROLES:
        return jsonify({"error": "Unauthorized"}), 403
    client_id = request.args.get('client_id')
    if not client_id:
        return jsonify({"error": "client_id required"}), 400
    data = WorkLog.get_client_summary(
        int(client_id),
        request.args.get('start_date'),
        request.args.get('end_date')
    )
    # Serialize nested dicts
    for key in ['by_department', 'by_task', 'by_employee', 'daily_timeline']:
        data[key] = _serialize(data[key])
    if data.get('overall'):
        data['overall'] = _serialize([data['overall']])[0]
    return jsonify(data), 200


@reports_bp.route('/reports/department', methods=['GET'])
@jwt_required()
def report_department():
    claims = get_jwt()
    if claims['role'] not in ADMIN_ROLES:
        return jsonify({"error": "Unauthorized"}), 403
    rows = WorkLog.get_department_summary(
        department=request.args.get('department'),
        start_date=request.args.get('start_date'),
        end_date=request.args.get('end_date')
    )
    return jsonify(_serialize(rows)), 200


@reports_bp.route('/reports/employee', methods=['GET'])
@jwt_required()
def report_employee():
    claims = get_jwt()
    if claims['role'] not in ADMIN_ROLES:
        return jsonify({"error": "Unauthorized"}), 403
    employee_id = request.args.get('employee_id')
    if not employee_id:
        return jsonify({"error": "employee_id required"}), 400
    rows = WorkLog.get_employee_summary(
        int(employee_id),
        request.args.get('start_date'),
        request.args.get('end_date')
    )
    return jsonify(_serialize(rows)), 200


@reports_bp.route('/reports/full', methods=['GET'])
@jwt_required()
def report_full():
    claims = get_jwt()
    if claims['role'] not in ADMIN_ROLES:
        return jsonify({"error": "Unauthorized"}), 403
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    if not start_date or not end_date:
        return jsonify({"error": "start_date and end_date required"}), 400
    rows = WorkLog.get_full_company_summary(start_date, end_date)
    return jsonify(_serialize(rows)), 200


@reports_bp.route('/reports/export-csv', methods=['POST'])
@jwt_required()
def export_csv():
    claims = get_jwt()
    if claims['role'] not in ADMIN_ROLES:
        return jsonify({"error": "Unauthorized"}), 403
    data = request.json
    rows = data.get('rows', [])
    report_type = data.get('report_type', 'report')
    if not rows:
        return jsonify({"error": "No data"}), 400

    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=rows[0].keys())
    writer.writeheader()
    writer.writerows(rows)

    response = make_response(output.getvalue())
    response.headers['Content-Type'] = 'text/csv'
    response.headers['Content-Disposition'] = f'attachment; filename={report_type}_{datetime.now().strftime("%Y%m%d")}.csv'
    return response


@reports_bp.route('/reports/generate-pdf', methods=['POST'])
@jwt_required()
def generate_pdf():
    claims = get_jwt()
    user_id = int(get_jwt_identity())
    if claims['role'] not in ADMIN_ROLES:
        return jsonify({"error": "Unauthorized"}), 403

    data = request.json
    rows = data.get('rows', [])
    report_type = data.get('report_type', 'Work Summary Report')
    start_date = data.get('start_date', '')
    end_date = data.get('end_date', '')
    columns = data.get('columns', [])
    totals = data.get('totals', {})

    generator = User.get_by_id(user_id)
    generator_name = generator['name'] if generator else 'Admin'

    COL_LABELS = {
        'user_name': 'Employee', 'department_name': 'Department', 'task_title': 'Task',
        'log_date': 'Date', 'work_date': 'Date', 'start_time': 'Start', 'end_time': 'End',
        'duration_minutes': 'Duration', 'hours_worked': 'Hours', 'status': 'Status',
        'company_name': 'Client', 'team_name': 'Team', 'employee_name': 'Employee',
        'department': 'Department',
    }

    def img_path(filename):
        return os.path.normpath(
            os.path.join(os.path.dirname(__file__), '..', '..', '..', 'frontend', filename)
        )

    top_path    = img_path('letterpadtop.png')
    bottom_path = img_path('letterpadbottom.png')

    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Image, Table, TableStyle, Spacer, Paragraph
    from reportlab.lib.styles import ParagraphStyle
    from PIL import Image as PILImage

    FONT_NORMAL, FONT_BOLD = _get_pdf_fonts()

    PAGE_W, PAGE_H = A4

    def img_dims(path):
        """Return (width_pts, height_pts) scaled to full page width."""
        pil = PILImage.open(path)
        iw, ih = pil.size
        h = PAGE_W * (ih / iw)
        return PAGE_W, h

    top_w,    top_h    = img_dims(top_path)    if os.path.exists(top_path)    else (0, 0)
    bottom_w, bottom_h = img_dims(bottom_path) if os.path.exists(bottom_path) else (0, 0)

    # ── Build table data ──
    header_row = [COL_LABELS.get(c, c.replace('_', ' ').title()) for c in columns]
    table_data = [header_row]
    for row in rows:
        table_data.append([str(row.get(c, '') or '') for c in columns])
    if totals:
        table_data.append([str(totals.get(c, '') or '') for c in columns])

    col_count = len(columns) or 1

    title_style  = ParagraphStyle('t', fontSize=13, textColor=colors.HexColor('#2563eb'),
                                   spaceAfter=2, fontName=FONT_BOLD)
    sub_style    = ParagraphStyle('s', fontSize=9,  textColor=colors.HexColor('#555555'),
                                   spaceAfter=4, fontName=FONT_NORMAL)
    footer_style = ParagraphStyle('f', fontSize=8,  textColor=colors.HexColor('#888888'),
                                   fontName=FONT_NORMAL)

    SIDE_PAD = 14 * mm
    inner_w  = PAGE_W - 2 * SIDE_PAD
    col_w    = inner_w / col_count

    tbl = Table(table_data, colWidths=[col_w] * col_count, repeatRows=1)
    tbl.setStyle(TableStyle([
        ('BACKGROUND',     (0, 0),  (-1, 0),  colors.HexColor('#dbeafe')),
        ('TEXTCOLOR',      (0, 0),  (-1, 0),  colors.HexColor('#1e3a8a')),
        ('FONTNAME',       (0, 0),  (-1, 0),  FONT_BOLD),
        ('FONTSIZE',       (0, 0),  (-1, 0),  8),
        ('TOPPADDING',     (0, 0),  (-1, 0),  5),
        ('BOTTOMPADDING',  (0, 0),  (-1, 0),  5),
        ('FONTNAME',       (0, 1),  (-1, -2), FONT_NORMAL),
        ('FONTSIZE',       (0, 1),  (-1, -2), 7.5),
        ('ROWBACKGROUNDS', (0, 1),  (-1, -2), [colors.white, colors.HexColor('#f9fafb')]),
        ('TOPPADDING',     (0, 1),  (-1, -2), 4),
        ('BOTTOMPADDING',  (0, 1),  (-1, -2), 4),
        ('BACKGROUND',     (0, -1), (-1, -1), colors.HexColor('#f0f0f0')),
        ('FONTNAME',       (0, -1), (-1, -1), FONT_BOLD),
        ('FONTSIZE',       (0, -1), (-1, -1), 8),
        ('GRID',           (0, 0),  (-1, -1), 0.4, colors.HexColor('#cccccc')),
        ('VALIGN',         (0, 0),  (-1, -1), 'MIDDLE'),
    ]))

    def build_elements():
        elems = [
            Paragraph(report_type, title_style),
            Paragraph(f'Date Range: {start_date}  to  {end_date}', sub_style),
            Spacer(1, 4 * mm),
            tbl,
            Spacer(1, 5 * mm),
            Paragraph(
                f'Generated by: {generator_name} &nbsp;&nbsp;&nbsp; '
                f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}',
                footer_style
            ),
        ]
        # Wrap with side padding
        wrapper = Table([[ elems ]], colWidths=[inner_w])
        wrapper.setStyle(TableStyle([
            ('LEFTPADDING',   (0, 0), (-1, -1), SIDE_PAD),
            ('RIGHTPADDING',  (0, 0), (-1, -1), SIDE_PAD),
            ('TOPPADDING',    (0, 0), (-1, -1), 5 * mm),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4 * mm),
        ]))
        return [wrapper]

    # Top image drawn on every page via callback
    def on_page(canvas, doc):
        canvas.saveState()
        if os.path.exists(top_path) and top_h > 0:
            canvas.drawImage(top_path, 0, PAGE_H - top_h,
                             width=PAGE_W, height=top_h,
                             preserveAspectRatio=True, mask='auto')
        canvas.restoreState()

    # Bottom image as last flowable — always at end of content
    def all_elements():
        elems = build_elements()
        if os.path.exists(bottom_path) and bottom_h > 0:
            elems.append(Spacer(1, 4 * mm))
            elems.append(Image(bottom_path, width=PAGE_W, height=bottom_h))
        return elems

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=0, rightMargin=0,
        topMargin=top_h, bottomMargin=0,
    )
    doc.build(all_elements(), onFirstPage=on_page, onLaterPages=on_page)

    buf.seek(0)
    response = make_response(buf.read())
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = (
        f'attachment; filename={report_type.replace(" ", "_")}_{datetime.now().strftime("%Y%m%d")}.pdf'
    )
    return response
@reports_bp.route('/reports/salary-report', methods=['GET'])
@jwt_required()
def salary_report_pdf():
    claims = get_jwt()
    if claims['role'] not in ADMIN_ROLES:
        return jsonify({"error": "Unauthorized"}), 403

    month  = request.args.get('month', datetime.now().month)
    year   = request.args.get('year', datetime.now().year)
    org_id = claims.get('organisation_id')

    from app.models.salary import Salary
    from app.utils.database import get_db_connection
    conn   = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    org_f  = "AND organisation_id = %s" if org_id is not None else ""
    params = ([org_id] if org_id is not None else [])
    cursor.execute(f"SELECT id, name, role FROM users WHERE role != 'client' {org_f} ORDER BY name", params)
    employees = cursor.fetchall()
    cursor.close(); conn.close()

    rows = []
    total_payout = 0
    for emp in employees:
        calc = Salary.calculate_expected_salary(emp['id'], int(month), int(year))
        row = {
            'name': emp['name'],
            'role': emp['role'].replace('_', ' ').title(),
            'working_days': calc['total_working_days'],
            'present': calc['present_days'],
            'leaves': calc['leaves'],
            'base': f"₹{calc['base_salary']:,.2f}" if 'base_salary' in calc else "₹0.00",
            'allowance': f"₹{calc['allowance']:,.2f}" if 'allowance' in calc else "₹0.00",
            'total': calc['expected_salary']
        }
        # Simplify display
        row['base_val'] = calc.get('base_salary', 0)
        row['allow_val'] = calc.get('allowance', 0)
        rows.append(row)
        total_payout += calc['expected_salary']

    # ── PDF Generation ──
    import io, os
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Image, Table, TableStyle, Spacer, Paragraph
    from reportlab.lib.styles import ParagraphStyle
    from PIL import Image as PILImage

    buf = io.BytesIO()
    PAGE_W, PAGE_H = A4

    FONT_NORMAL, FONT_BOLD = _get_pdf_fonts()
    
    def img_path(filename):
        return os.path.normpath(os.path.join(os.path.dirname(__file__), '..', '..', '..', 'frontend', filename))
    
    top_path = img_path('letterpadtop.png')
    bottom_path = img_path('letterpadbottom.png')

    def img_dims(path):
        if not os.path.exists(path): return 0, 0
        pil = PILImage.open(path); iw, ih = pil.size
        return PAGE_W, PAGE_W * (ih / iw)

    top_w, top_h = img_dims(top_path)
    
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=15*mm, rightMargin=15*mm, topMargin=top_h + 5*mm, bottomMargin=20*mm)
    
    title_style = ParagraphStyle('t', fontSize=16, textColor=colors.HexColor('#1e40af'), spaceAfter=10, fontName=FONT_BOLD, alignment=1)
    sub_style = ParagraphStyle('s', fontSize=10, textColor=colors.HexColor('#4b5563'), spaceAfter=15, fontName=FONT_NORMAL, alignment=1)
    
    months = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
    report_title = f"Salary Sheet - {months[int(month)-1]} {year}"
    
    # Table headers
    headers = ['Employee', 'Role', 'Days', 'Pres.', 'Leaves', 'Base', 'Allow.', 'Total']
    table_data = [headers]
    for r in rows:
        table_data.append([
            r['name'], r['role'], r['working_days'], r['present'], r['leaves'],
            f"₹{r['base_val']:,.0f}", f"₹{r['allow_val']:,.0f}", f"₹{r['total']:,.2f}"
        ])
    
    # Totals row
    table_data.append(['', '', '', '', '', '', 'GRAND TOTAL:', f"₹{total_payout:,.2f}"])

    tbl = Table(table_data, colWidths=[35*mm, 25*mm, 12*mm, 12*mm, 15*mm, 22*mm, 20*mm, 28*mm], repeatRows=1)
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#eff6ff')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#1e40af')),
        ('FONTNAME', (0,0), (-1,0), FONT_BOLD),
        ('FONTSIZE', (0,0), (-1,0), 9),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-2), 0.5, colors.grey),
        ('FONTNAME', (0,1), (-1,-2), FONT_NORMAL),
        ('FONTSIZE', (0,1), (-1,-2), 8),
        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#f8fafc')),
        ('FONTNAME', (0,-1), (-1,-1), FONT_BOLD),
        ('SPAN', (0,-1), (6,-1)),
        ('ALIGN', (0,-1), (-1,-1), 'RIGHT'),
    ]))

    def on_page(canvas, doc):
        canvas.saveState()
        if os.path.exists(top_path):
            canvas.drawImage(top_path, 0, PAGE_H - top_h, width=PAGE_W, height=top_h, preserveAspectRatio=True, mask='auto')
        if os.path.exists(bottom_path):
            _, bh = img_dims(bottom_path)
            canvas.drawImage(bottom_path, 0, 0, width=PAGE_W, height=bh, preserveAspectRatio=True, mask='auto')
        canvas.restoreState()

    elements = [
        Paragraph(report_title, title_style),
        Paragraph(f"Generated on {datetime.now().strftime('%d %B %Y, %H:%M')}", sub_style),
        tbl
    ]
    
    doc.build(elements, onFirstPage=on_page, onLaterPages=on_page)
    buf.seek(0)
    
    response = make_response(buf.read())
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'attachment; filename=Salary_Report_{month}_{year}.pdf'
    return response

@reports_bp.route('/reports/salary-report-docx', methods=['GET'])
@jwt_required()
def salary_report_docx():
    claims = get_jwt()
    if claims['role'] not in ADMIN_ROLES:
        return jsonify({"error": "Unauthorized"}), 403

    month  = request.args.get('month', datetime.now().month)
    year   = request.args.get('year', datetime.now().year)
    org_id = claims.get('organisation_id')

    from app.models.salary import Salary
    from app.utils.database import get_db_connection
    conn   = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    org_f  = "AND organisation_id = %s" if org_id is not None else ""
    params = ([org_id] if org_id is not None else [])
    cursor.execute(f"SELECT id, name, role FROM users WHERE role != 'client' {org_f} ORDER BY name", params)
    employees = cursor.fetchall()
    cursor.close(); conn.close()

    months = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
    report_title = f"Salary Sheet - {months[int(month)-1]} {year}"

    import io
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    doc.add_heading(report_title, 0)
    
    p = doc.add_paragraph(f"Generated on {datetime.now().strftime('%d %B %Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    headers = ['Employee', 'Role', 'Days', 'Pres.', 'Leaves', 'Base', 'Allow.', 'Total']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        # Make bold
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    total_payout = 0
    for emp in employees:
        calc = Salary.calculate_expected_salary(emp['id'], int(month), int(year))
        row_cells = table.add_row().cells
        row_cells[0].text = emp['name']
        row_cells[1].text = emp['role'].replace('_', ' ').title()
        row_cells[2].text = str(calc['total_working_days'])
        row_cells[3].text = str(calc['present_days'])
        row_cells[4].text = str(calc['leaves'])
        row_cells[5].text = f"₹{calc.get('base_salary', 0):,.0f}"
        row_cells[6].text = f"₹{calc.get('allowance', 0):,.0f}"
        row_cells[7].text = f"₹{calc['expected_salary']:,.2f}"
        total_payout += calc['expected_salary']

    # Totals row
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[6])
    row_cells[0].text = "GRAND TOTAL:"
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[7].text = f"₹{total_payout:,.2f}"
    row_cells[7].paragraphs[0].runs[0].font.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    response = make_response(buf.read())
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response.headers['Content-Disposition'] = f'attachment; filename=Salary_Report_{month}_{year}.docx'
    return response
